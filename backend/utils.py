import os
from pathlib import Path
import PyPDF2
import pdfplumber
from docx import Document
import pandas as pd

def load_personality():
    """Load and concatenate content from resume files and personal info for AI system prompt."""
    all_content = []
    
    # Define the specific files to read
    files_to_read = [
        "me.txt",
        "Nikhil_Suresh_Resume.docx", 
        "Nikhil_Suresh_Resume_Updated.pdf"
    ]
    
    for file_path in files_to_read:
        if not os.path.exists(file_path):
            continue
            
        try:
            file_extension = Path(file_path).suffix.lower()
            content = ""
            
            # Read TXT file
            if file_extension == '.txt':
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read().strip()
            
            # Read PDF file
            elif file_extension == '.pdf':
                pdf_content_parts = []
                
                # Extract text content
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text_content = []
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(page_text.strip())
                    
                    if text_content:
                        pdf_content_parts.extend(text_content)
                
                # Extract tables using pdfplumber
                try:
                    with pdfplumber.open(file_path) as pdf:
                        for page in pdf.pages:
                            tables = page.extract_tables()
                            if tables:
                                for table in tables:
                                    if table and any(any(cell for cell in row if cell) for row in table):
                                        try:
                                            # Convert table to readable format
                                            table_text = []
                                            for row in table:
                                                if row and any(cell for cell in row if cell):
                                                    cleaned_row = [str(cell).strip() if cell else "" for cell in row]
                                                    table_text.append(" | ".join(cleaned_row))
                                            
                                            if table_text:
                                                pdf_content_parts.append("\n".join(table_text))
                                        except Exception:
                                            # Fallback for malformed tables
                                            for row in table:
                                                if row:
                                                    pdf_content_parts.append(" | ".join(str(cell) for cell in row if cell))
                
                except Exception:
                    pass  # Silently handle table extraction errors
                
                content = "\n".join(pdf_content_parts)
            
            # Read DOCX file
            elif file_extension == '.docx':
                doc = Document(file_path)
                docx_content_parts = []
                
                # Read paragraphs
                for para in doc.paragraphs:
                    if para.text.strip():
                        docx_content_parts.append(para.text.strip())
                
                # Read tables
                if doc.tables:
                    for table in doc.tables:
                        table_data = []
                        for row in table.rows:
                            row_data = [cell.text.strip() for cell in row.cells]
                            if any(cell for cell in row_data if cell):
                                table_data.append(row_data)
                        
                        if table_data:
                            try:
                                # Convert table to readable text
                                table_text = []
                                for row in table_data:
                                    cleaned_row = [cell if cell else "" for cell in row]
                                    table_text.append(" | ".join(cleaned_row))
                                
                                if table_text:
                                    docx_content_parts.append("\n".join(table_text))
                            except Exception:
                                # Fallback
                                for row in table_data:
                                    docx_content_parts.append(" | ".join(str(cell) for cell in row if cell))
                
                content = "\n".join(docx_content_parts)
            
            # Add content if not empty
            if content and content.strip():
                all_content.append(content.strip())
                
        except Exception:
            continue  # Silently handle file reading errors
    
    # Create a comprehensive system prompt
    if all_content:
        combined_content = "\n\n".join(all_content)
        
        # Format as system prompt
        system_prompt = f"""You are Nikhil Suresh's AI digital twin. You should respond as if you are Nikhil himself, using first person ("I", "my", "me"). You have access to all of Nikhil's professional experience, skills, projects, and personal information.

Here is all the information about Nikhil:

{combined_content}

Instructions:
- Always respond in first person as Nikhil Suresh
- Use the information provided to answer questions about experience, skills, projects, education, etc.
- Be conversational and professional
- If asked about something not in the provided information, politely indicate you don't have that specific information
- Maintain Nikhil's personality and professional tone
- Feel free to elaborate on technical topics within your expertise"""
        
        return system_prompt.strip()
    else:
        return "You are Nikhil Suresh's AI digital twin. Please respond as Nikhil in first person, but I don't have access to detailed information at the moment."